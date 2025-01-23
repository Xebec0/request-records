from django.shortcuts import render, get_object_or_404
from django.http import FileResponse, Http404
from docx2pdf import convert
from docx import Document
import os
from django.conf import settings
from ..models import Profile, ReportTemplate, Purpose
from datetime import datetime

def index(request):
    templates = ReportTemplate.objects.all()
    purposes = Purpose.objects.filter(active=True)
    return render(request, 'user/reports.html', {'templates': templates, 'purposes': purposes})

def generate_pdf(request, template_id):
    # Get template and user profile
    purpose_id = request.GET.get('purpose')
    if not purpose_id:
        raise Http404("Purpose is required")

    try:
        purpose = Purpose.objects.get(id=purpose_id)

        template = get_object_or_404(ReportTemplate, id=template_id)
        profile = get_object_or_404(Profile, user=request.user)
        
        # Setup directories
        generated_dir = os.path.join(settings.MEDIA_ROOT, 'reports', 'generated')
        os.makedirs(generated_dir, exist_ok=True)
        
        # Use template name for files
        filename_base = template.name
        docx_path = os.path.join(generated_dir, f"{filename_base}.docx")
        pdf_path = os.path.join(generated_dir, f"{filename_base}.pdf")
        
        # Define field mappings for template
        field_mapping = {
            'first_name': profile.first_name,
            'last_name': profile.last_name,
            'middle_name': profile.middle_name,
            'full_name': f"{profile.first_name} {profile.middle_name} {profile.last_name}",
            'contact_no': str(profile.contact_no),
            'entry_year_from': str(profile.entry_year_from),
            'entry_year_to': str(profile.entry_year_to),
            'course': profile.course.description,
            'student_number': profile.user.student_number,
            'email': profile.user.email,
            'current_date': datetime.now().strftime('%B %d, %Y'),
            'purpose': purpose.description
        }

        # Process document
        doc = Document(template.template_file.path)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for field_name, field_value in field_mapping.items():
                placeholder = f'{{{{{field_name}}}}}'
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, field_value)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for field_name, field_value in field_mapping.items():
                            placeholder = f'{{{{{field_name}}}}}'
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, field_value)
        
        # Save and convert
        doc.save(docx_path)
        convert(docx_path, pdf_path)
        
        # Clean up DOCX
        if os.path.exists(docx_path):
            os.remove(docx_path)

        # Stream PDF file for display in browser
        response = FileResponse(
            open(pdf_path, 'rb'),
            content_type='application/pdf',
            as_attachment=False,  # Set to False for in-browser display
            filename=f"{filename_base}.pdf"
        )
        
        # Add file cleanup after streaming
        response._resource_closers.append(lambda: os.remove(pdf_path))
        
        return response
                
    except Exception as e:
        raise Http404(f"Error generating PDF: {str(e)}")
